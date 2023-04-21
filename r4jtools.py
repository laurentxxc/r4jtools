#!/usr/bin/env python
import argparse
import getpass
import yaml
import json
import requests
import sys

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

JIRA_PATH = '/rest/api/latest'
R4J_PATH = '/rest/com.easesolutions.jira.plugins.requirements/2.0'
INSERTION_POINT_STYLE = 'Insertion Point'
REQUIREMENTS_TAG = '<Requirements>'


# Style
TABLE_STYLE = 'Requirement_list'
REQ_TITLE_STYLE = 'Requirement Title'
REQ_STYLE = 'Requirement'


TEST = False

# Tools for exporting R4J requirements in word doc

def log(*args):
    print(*args, file=sys.stderr)

class JiraSearch(object):
    """ This factory will create the actual method used to fetch issues from JIRA. This is really just a closure that
        saves us having to pass a bunch of parameters all over the place all the time. """

    __base_url = None

    def __init__(self, url, auth, no_verify_ssl):
        self.__base_url = url
        self.url = url
        self.auth = auth

        self.no_verify_ssl = no_verify_ssl
        self.fields = ','.join(['key', 'summary', 'status', 'description', 'issuetype', 'issuelinks', 'subtasks','updated'])

    def get(self, uri, params={}):
        headers = {'Content-Type' : 'application/json'}
        url = self.url + uri

        if isinstance(self.auth, str):
            return requests.get(url, params=params, cookies={'JSESSIONID': self.auth}, headers=headers, verify=self.no_verify_ssl)
        else:
            return requests.get(url, params=params, auth=self.auth, headers=headers, verify=(not self.no_verify_ssl))

    def get_issue(self, key):
        """ Given an issue key (i.e. JRA-9) return the JSON representation of it. This is the only place where we deal
            with JIRA's REST API. """
        log('Fetching ' + key)
        # we need to expand subtasks and links since that's what we care about here.
        response = self.get(f'{JIRA_PATH}/issue/{key}', params={'fields': self.fields})
        response.raise_for_status()
        return response.json()

    def query(self, query):
        log('Querying ' + query)
        response = self.get(f'{JIRA_PATH}/search', params={'jql': query, 'fields': self.fields})
        content = response.json()
        return content['issues']

    def list_ids(self, query):
        log('Querying ' + query)
        response = self.get(f'{JIRA_PATH}/search', params={'jql': query, 'fields': 'key', 'maxResults': 500})
        return [issue["key"] for issue in response.json()["issues"]]

    def get_issue_uri(self, issue_key):
        return self.__base_url + '/browse/' + issue_key
    
    def get_requirements_tree(self, projkey):
        if TEST:
            with open('reqSampleFolders.json') as testfile:
                return json.load(testfile)
        else:
            response = self.get(f'{R4J_PATH}/projects/{projkey}/folders', params={'plugin':'r4j'})
            return response.json()
    
    def get_requirement(self, reqKey):
        return self.get_issue(reqKey)

def parse_args():
    parser = argparse.ArgumentParser(description='programm_description')
    parser.add_argument('-c, --config', dest='config', default='config.yaml', help='Configuration file. By default config.yaml file is used')
    parser.add_argument('-u', '--user', dest='user', default=None, help='Username to access JIRA')
    parser.add_argument('-p', '--password', dest='password', default=None, help='Password to access JIRA')
    return parser.parse_args()


def printFolderNames(indent: str, jsonData):
    print(indent + jsonData['name'])
    for subfolder in jsonData['folders']:
        printFolderNames(f"{indent}\t",subfolder)


def walkFolderReq(lvl:int, folderJson):
    # TODO: how to handle case where requirement folder is empty or there is no requirement inside.
    # - Should we export folder as a requirement section? 
    # TODO: export in word document
    # TODO: export additional information
    # - Jira hyperling
    # - requirement description (need specific Jira request)
    # TODO: Apply filtering (product)

    indent = ' '*lvl

    def exportFolderInfo():
        print(f"{indent}{folderJson['name']}: {folderJson['description']}")

    def exportReqs():
        for issue in folderJson['issues']:
            req = issue['data']
            print(f"{indent}--> {req['key']}: {req['fields']['summary']}")
            print(f"{indent}\t{req['fields']['description']}")

    exportFolderInfo()
    exportReqs()
    for subfolder in folderJson['folders']:
        walkFolderReq(lvl+1, subfolder)


def exportR4JRequirements(doc, r4jfolders, insertPoint, search: JiraSearch):
    def recExport(lvl:int, folderJson):
        # TODO: how to handle case where requirement folder is empty or there is no requirement inside.
        # - Should we export folder as a requirement section? 
        # TODO: export in word document
        # TODO: export additional information
        # - Jira hyperling
        # - requirement description (need specific Jira request)
        # TODO: Apply filtering (product)


        # Workaround to move a table as no API exists (see https://github.com/python-openxml/python-docx/issues/156)
        def move_table_next(table: Table, destObj):
            tbl = table._tbl
            if isinstance(destObj,Paragraph):
                dest = destObj._p
            if isinstance(destObj,Table):
                dest = destObj._tbl
            dest.addnext(tbl)

        def move_table_before(table: Table, paragraph: Paragraph):
            tbl, p = table._tbl, paragraph._p
            p.insertbefore(tbl)


        def exportFolderInfo() -> Paragraph :
            title = folderJson['name']
            desc = folderJson['description']
            style = f"Heading {lvl+1}"
            insertPoint.insert_paragraph_before(title, doc.styles[style])
            return insertPoint.insert_paragraph_before(desc,doc.styles['Normal'])

        def exportReqs(prevParagraph):
            ref = prevParagraph
            for issue in folderJson['issues']:
                # create table
                reqTable: Table = doc.add_table(2,1) 
                reqTable.style = doc.styles[TABLE_STYLE]

                req = issue['data']
                reqKey = req['key']
                reqTitle = f"{reqKey}: {req['fields']['summary']}"

                ## request require details
                reqJson = search.get_requirement(reqKey)
                reqDesc = reqJson['fields']['description']
                
                # req title
                p =  reqTable.cell(0,0).paragraphs[0]
                p.add_run(reqTitle)
                p.style = doc.styles[REQ_TITLE_STYLE]
               
               # req description
                p = reqTable.cell(1,0).paragraphs[0]
                p.add_run(reqDesc)
                p.style = doc.styles[REQ_STYLE]

                move_table_next(reqTable, ref)
                #move_table_before(reqTable, insertPoint)

                # add empty paragraph before next req
                ref = insertPoint.insert_paragraph_before('',doc.styles['Normal'])


        folderPara = exportFolderInfo()
        exportReqs(folderPara)
        for subfolder in folderJson['folders']:
            recExport(lvl+1, subfolder)
    
    
    recExport(1,r4jfolders)



def findInsertionPoint(doc, name) -> Paragraph:
    for paragraph in doc.paragraphs:
        if (paragraph.style.name == INSERTION_POINT_STYLE) and (name in paragraph.text):
            return paragraph
    return None

def main():
    options = parse_args()
    
    # 3 steps:
    # - connect Jira database specified in config
    ## get yaml measurement scenario
    with open(options.config, 'r') as file:
        config = yaml.safe_load(file)
    #print(config)

    ## get Jira credential
    user = options.user if options.user is not None \
                else input('Username: ')
    password = options.password if options.password is not None \
                else getpass.getpass('Password: ')
    auth = (user, password)
    
    # - get requirements from the database
    jira = JiraSearch("https://%s" % config['r4j']['server'], auth, False)
    treeReqJson = jira.get_requirements_tree(projkey=config['r4j']['projkey'])

    #print(treeReqJson)
    #printFolderNames("", treeReqJson)

    # - Export requirement in word doc
    exportDoc = Document(config['template'])
    insertPoint = findInsertionPoint(exportDoc,REQUIREMENTS_TAG)
    exportR4JRequirements(exportDoc, treeReqJson, insertPoint, jira)
    exportDoc.save(config['output'])

    # requirement presentation idea:
    # folder has Heading N+1 style (N = folder level)
    # 1 column table with first line
    # # Req ID + Name as specific heading style (is it possible to have it in table header)
    # # description
    # 
    #walkFolderReq(0,treeReqJson)

if __name__ == '__main__':
    main()